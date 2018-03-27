from docx import Document
import random
import sys

vulns = []
with open('vulns.txt') as fp:
    for line in fp:
        parts = line.split(',', 1)
        vulns.append((parts[0].strip(), parts[1].strip()))

rand = random.SystemRandom()
rand.shuffle(vulns)
vuln_id = 0

doc = Document("template.docx")
for table in doc.tables:
    for cell in table._cells:
        title, description = vulns[vuln_id]
        vuln_id += 1
        for para in cell.paragraphs:
            if para.text == "Vuln":
                para.text = title
            elif para.text == "description":
                para.text = description

try:
    outfile = sys.argv[1]
    if outfile == "-":
        outfile = sys.stdout.buffer
    doc.save(outfile)
except IndexError:
    print("Usage: vulnbingo.py [outfile.docx]")
